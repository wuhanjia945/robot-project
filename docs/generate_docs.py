from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

def set_cell_shading(cell, color):
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0x7a)
    return heading

def create_version_record():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)

    title = doc.add_heading('一起动手来手搓机器人 - 版本记录', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0x7a)

    doc.add_paragraph('')

    table = doc.add_table(rows=2, cols=6)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['版本号', '日期', '修改人', '模块', '修改内容', '备注']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    row1_data = ['v1.0.0', '2026-05-25', '系统开发组', '全模块', '初始版本发布：\n1. 需求采集引擎（3轮结构化问卷）\n2. 方案生成器（架构+BOM+算法选型）\n3. 算法适配器（22个算法定义）\n4. 矛盾检测引擎（7类矛盾检测）\n5. 硬件方案库（26项硬件数据）\n6. 技术栈知识库（24项技术栈）\n7. React前端界面\n8. Docker容器化部署', 'MVP版本']
    for i, data in enumerate(row1_data):
        cell = table.rows[1].cells[i]
        cell.text = data
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    doc.add_paragraph('')
    add_heading_styled(doc, '版本规划', level=1)

    plans = [
        ('v1.1.0', 'LLM驱动的自由对话需求采集，替代固定问卷'),
        ('v1.2.0', 'RAG增强的方案推荐，接入向量数据库'),
        ('v1.3.0', '代码脚手架生成，基于Jinja2模板'),
        ('v2.0.0', '强化学习运动策略生成，仿真环境集成'),
        ('v2.1.0', '多机器人协同方案支持'),
        ('v3.0.0', '完整IDE集成，在线代码编辑与部署'),
    ]

    plan_table = doc.add_table(rows=len(plans)+1, cols=3)
    plan_table.style = 'Light Grid Accent 1'
    plan_headers = ['版本号', '核心功能', '状态']
    for i, h in enumerate(plan_headers):
        plan_table.rows[0].cells[i].text = h
        for p in plan_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    for idx, (ver, desc) in enumerate(plans):
        plan_table.rows[idx+1].cells[0].text = ver
        plan_table.rows[idx+1].cells[1].text = desc
        plan_table.rows[idx+1].cells[2].text = '计划中'

    filepath = os.path.join(OUTPUT_DIR, '版本记录.docx')
    doc.save(filepath)
    print(f"Created: {filepath}")

def create_project_manual():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)

    title = doc.add_heading('一起动手来手搓机器人 - 项目说明书', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0x7a)

    add_heading_styled(doc, '1. 项目概述', level=1)
    doc.add_paragraph('一起动手来手搓机器人是一套面向机器人开发者的全流程辅助平台，让每个人都能手搓一台机器人。提供从需求分析到技术方案生成、算法适配、硬件选型的一站式服务。系统支持两种使用模式：全栈方案模式（从零开始设计机器人）和算法供给模式（已有硬件，获取配套算法）。')

    add_heading_styled(doc, '2. 系统架构', level=1)
    doc.add_paragraph('系统采用前后端分离架构：')
    doc.add_paragraph('后端：Python FastAPI，提供RESTful API', style='List Bullet')
    doc.add_paragraph('前端：React + TypeScript + Vite', style='List Bullet')
    doc.add_paragraph('数据存储：YAML结构化知识库', style='List Bullet')
    doc.add_paragraph('部署：Docker + docker-compose容器化', style='List Bullet')

    add_heading_styled(doc, '3. 配置参数说明', level=1)

    add_heading_styled(doc, '3.1 后端配置', level=2)
    config_table = doc.add_table(rows=7, cols=4)
    config_table.style = 'Light Grid Accent 1'
    config_headers = ['参数名', '默认值', '类型', '说明']
    for i, h in enumerate(config_headers):
        config_table.rows[0].cells[i].text = h
        for p in config_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    configs = [
        ('APP_NAME', '一起动手来手搓机器人', 'str', '应用名称'),
        ('APP_VERSION', '1.0.0', 'str', '应用版本'),
        ('DATA_DIR', 'backend/data', 'Path', '知识库数据目录'),
        ('HOST', '0.0.0.0', 'str', '服务监听地址'),
        ('PORT', '8000', 'int', '服务监听端口'),
        ('DEBUG', 'True', 'bool', '调试模式'),
    ]
    for idx, (name, default, type_, desc) in enumerate(configs):
        config_table.rows[idx+1].cells[0].text = name
        config_table.rows[idx+1].cells[1].text = default
        config_table.rows[idx+1].cells[2].text = type_
        config_table.rows[idx+1].cells[3].text = desc

    add_heading_styled(doc, '3.2 环境变量', level=2)
    doc.add_paragraph('所有配置参数均支持通过环境变量覆盖，前缀为 APP_，例如：APP_PORT=9000')

    add_heading_styled(doc, '4. 知识库扩展指南', level=1)

    add_heading_styled(doc, '4.1 技术栈扩展', level=2)
    doc.add_paragraph('在 backend/data/tech_stack/ 目录下添加或修改YAML文件即可扩展技术栈知识库。文件格式：')
    doc.add_paragraph('category: 分类名称', style='List Bullet')
    doc.add_paragraph('items: 技术栈条目列表', style='List Bullet')
    doc.add_paragraph('每个条目包含：name, version, purpose, language, license, compatibility, tags, description, url, difficulty', style='List Bullet')

    add_heading_styled(doc, '4.2 硬件方案扩展', level=2)
    doc.add_paragraph('在 backend/data/hardware/ 目录下添加或修改YAML文件。关键文件：')
    doc.add_paragraph('controllers.yaml - 主控制器', style='List Bullet')
    doc.add_paragraph('actuators.yaml - 执行器', style='List Bullet')
    doc.add_paragraph('sensors.yaml - 传感器', style='List Bullet')
    doc.add_paragraph('robot_kits.yaml - 完整机器人方案', style='List Bullet')
    doc.add_paragraph('power_communication.yaml - 电源与通信', style='List Bullet')

    add_heading_styled(doc, '4.3 算法库扩展', level=2)
    doc.add_paragraph('在 backend/data/algorithms/ 目录下添加或修改YAML文件。每个算法需包含：')
    doc.add_paragraph('name, level, applicable_robots, required_sensors, required_compute, description, difficulty, parameters, input_spec, output_spec, code_template', style='List Bullet')

    add_heading_styled(doc, '5. API接口说明', level=1)

    api_groups = [
        ('需求采集', '/api/requirement', [
            ('POST /session', '创建需求采集会话'),
            ('GET /session/{id}/round/{n}', '获取第n轮问题'),
            ('POST /session/{id}/round/{n}', '提交第n轮答案'),
            ('GET /session/{id}', '获取会话状态'),
            ('POST /session/{id}/build', '构建需求对象'),
            ('POST /hardware-algorithm', '硬件算法适配请求'),
        ]),
        ('方案生成', '/api/solution', [
            ('POST /generate', '根据需求生成技术方案'),
        ]),
        ('算法库', '/api/algorithm', [
            ('GET /list', '获取算法列表（支持category/level/robot_type筛选）'),
            ('GET /search', '搜索算法'),
            ('POST /adapt', '根据硬件信息适配算法'),
        ]),
        ('硬件方案', '/api/hardware', [
            ('GET /list', '获取硬件列表'),
            ('GET /categories', '获取分类列表'),
            ('GET /search', '搜索硬件'),
            ('GET /kits', '获取完整方案'),
            ('GET /budget', '按预算筛选'),
            ('GET /tech-stack', '获取技术栈'),
        ]),
        ('矛盾检测', '/api/validator', [
            ('POST /requirement', '检测需求矛盾'),
            ('POST /hardware-algorithm', '检测硬件与算法矛盾'),
        ]),
    ]

    for group_name, prefix, endpoints in api_groups:
        add_heading_styled(doc, f'5.{api_groups.index((group_name, prefix, endpoints))+1} {group_name} ({prefix})', level=2)
        ep_table = doc.add_table(rows=len(endpoints)+1, cols=2)
        ep_table.style = 'Light Grid Accent 1'
        ep_table.rows[0].cells[0].text = '接口'
        ep_table.rows[0].cells[1].text = '说明'
        for p in ep_table.rows[0].cells[0].paragraphs:
            for r in p.runs:
                r.font.bold = True
        for p in ep_table.rows[0].cells[1].paragraphs:
            for r in p.runs:
                r.font.bold = True
        for idx, (ep, desc) in enumerate(endpoints):
            ep_table.rows[idx+1].cells[0].text = ep
            ep_table.rows[idx+1].cells[1].text = desc

    add_heading_styled(doc, '6. 外部系统对接', level=1)
    doc.add_paragraph('系统通过RESTful API对外提供服务，其他智能体/系统可通过HTTP请求调用：')
    doc.add_paragraph('方案生成：POST /api/solution/generate，传入需求JSON，返回完整技术方案', style='List Bullet')
    doc.add_paragraph('算法适配：POST /api/algorithm/adapt，传入硬件信息，返回匹配算法', style='List Bullet')
    doc.add_paragraph('矛盾检测：POST /api/validator/requirement，传入需求，返回矛盾列表', style='List Bullet')
    doc.add_paragraph('健康检查：GET /health，返回系统状态', style='List Bullet')

    add_heading_styled(doc, '7. 部署指南', level=1)
    add_heading_styled(doc, '7.1 Docker部署（推荐）', level=2)
    doc.add_paragraph('docker-compose up -d')
    doc.add_paragraph('后端：http://localhost:8000')
    doc.add_paragraph('前端：http://localhost:3000')

    add_heading_styled(doc, '7.2 手动部署', level=2)
    doc.add_paragraph('后端：')
    doc.add_paragraph('cd backend && pip install -r requirements.txt && uvicorn app.main:app --host 0.0.0.0 --port 8000', style='List Bullet')
    doc.add_paragraph('前端：')
    doc.add_paragraph('cd frontend && npm install && npm run build && 使用nginx托管dist目录', style='List Bullet')

    add_heading_styled(doc, '8. 目录结构', level=1)
    structure = """backend/
├── app/
│   ├── api/          API路由层
│   ├── core/         核心引擎
│   ├── models/       数据模型
│   ├── schemas/      Pydantic模型
│   ├── config.py     配置管理
│   └── main.py       应用入口
├── data/
│   ├── tech_stack/   技术栈知识库
│   ├── hardware/     硬件方案库
│   ├── algorithms/   算法库
│   └── templates/    方案模板
├── tests/            测试用例
├── requirements.txt  Python依赖
└── Dockerfile
frontend/
├── src/
│   ├── pages/        页面组件
│   ├── components/   通用组件
│   ├── services/     API服务
│   ├── App.tsx       主应用
│   └── main.tsx      入口
├── package.json
└── Dockerfile
docker-compose.yaml   容器编排
docs/                 项目文档"""
    doc.add_paragraph(structure)

    filepath = os.path.join(OUTPUT_DIR, '项目说明书.docx')
    doc.save(filepath)
    print(f"Created: {filepath}")

if __name__ == '__main__':
    create_version_record()
    create_project_manual()
    print("All documents generated successfully!")
