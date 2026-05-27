from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

doc.add_heading('一起动手来手搓机器人 - 项目说明书', 0)

doc.add_heading('1. 项目概述', level=1)
doc.add_paragraph(
    '一起动手来手搓机器人是一个面向机器人开发者的智能辅助平台，让每个人都能手搓一台机器人。'
    '系统提供从需求分析到方案生成、'
    '算法适配、硬件选型的全流程决策支持。系统采用"预设优先+LLM兜底"架构，在保证响应速度的同时，'
    '通过大语言模型处理预设规则无法覆盖的复杂场景。'
)
doc.add_paragraph('核心功能：', style='List Bullet')
for item in [
    '需求采集：多轮对话式需求采集，自动检测需求矛盾',
    '方案生成：根据需求自动生成完整技术方案（架构/算法/BOM/成本估算）',
    '算法适配：22+算法库，按硬件平台自动适配推荐',
    '硬件方案：26+硬件组件，支持按预算/类型筛选完整方案',
    '市场方案：10种主流机器人方案，含技术框架/BOM/实现路径',
    '矛盾检测：自动检测需求矛盾和硬件-算法不匹配',
    '仿真学习：6种主流仿真软件对比与学习路径',
    '知识库：YAML结构化知识库，涵盖运动控制/视觉/SLAM/语音等6大技术栈',
]:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_heading('2. 系统架构', level=1)
doc.add_paragraph('系统采用前后端分离架构：')
doc.add_paragraph('后端：Python FastAPI + YAML知识库', style='List Bullet')
doc.add_paragraph('前端：React 18 + TypeScript + Vite', style='List Bullet')
doc.add_paragraph('部署：Docker Compose容器化', style='List Bullet')
doc.add_paragraph('LLM：支持OpenAI/DeepSeek/Ollama/vLLM/LocalAI', style='List Bullet')

doc.add_heading('2.1 后端架构', level=2)
doc.add_paragraph(
    '后端基于FastAPI框架，采用模块化设计，6大核心模块各自独立：'
)
add_table(doc,
    ['模块', '文件', '功能'],
    [
        ['需求引擎', 'core/requirement_engine.py', '多轮对话需求采集，矛盾检测'],
        ['方案生成器', 'core/solution_generator.py', '根据需求生成完整技术方案'],
        ['算法适配器', 'core/algorithm_adapter.py', '算法库管理，硬件适配推荐'],
        ['硬件目录', 'core/hardware_catalog.py', '硬件组件管理，预算筛选'],
        ['矛盾检测器', 'core/contradiction_detector.py', '需求矛盾/硬件算法不匹配检测'],
        ['知识库', 'core/knowledge_base.py', 'YAML知识库加载与查询'],
        ['LLM客户端', 'core/llm_client.py', '大语言模型API调用封装'],
    ]
)

doc.add_heading('2.2 前端架构', level=2)
add_table(doc,
    ['页面', '文件', '路由', '功能'],
    [
        ['首页', 'HomePage.tsx', '/', '系统概览与快速入口'],
        ['市场方案', 'MarketPage.tsx', '/market', '主流机器人方案浏览与详情'],
        ['需求采集', 'RequirementPage.tsx', '/requirement', '多轮对话式需求采集'],
        ['方案生成', 'SolutionPage.tsx', '/solution', '技术方案自动生成'],
        ['算法供给', 'AlgorithmPage.tsx', '/algorithm', '算法库浏览与适配'],
        ['硬件方案', 'HardwarePage.tsx', '/hardware', '硬件选型与方案配置'],
        ['仿真学习', 'SimulationPage.tsx', '/simulation', '仿真软件对比与学习'],
        ['反馈', 'FeedbackPage.tsx', '/feedback', '用户反馈收集'],
    ]
)

doc.add_heading('3. 配置参数说明', level=1)
doc.add_paragraph('所有配置通过环境变量设置，支持.env文件。环境变量前缀为APP_。')

doc.add_heading('3.1 基础配置', level=2)
add_table(doc,
    ['环境变量', '默认值', '说明'],
    [
        ['APP_HOST', '0.0.0.0', '服务监听地址'],
        ['APP_PORT', '8000', '服务监听端口'],
        ['APP_DEBUG', 'true', '调试模式（生产环境设为false）'],
    ]
)

doc.add_heading('3.2 LLM配置', level=2)
add_table(doc,
    ['环境变量', '默认值', '说明'],
    [
        ['APP_LLM_ENABLED', 'false', '是否启用LLM功能'],
        ['APP_LLM_API_KEY', '', 'LLM API密钥（本地模型不需要）'],
        ['APP_LLM_API_BASE', 'https://api.openai.com/v1', 'API基础URL'],
        ['APP_LLM_MODEL', 'gpt-4o-mini', '模型名称'],
        ['APP_LLM_TEMPERATURE', '0.3', '生成温度（0-1，越低越确定）'],
        ['APP_LLM_MAX_TOKENS', '4096', '最大生成Token数'],
        ['APP_LLM_TIMEOUT', '60', '请求超时时间（秒）'],
        ['APP_LLM_FALLBACK_ENABLED', 'true', 'LLM失败时是否回退到预设规则'],
    ]
)

doc.add_heading('3.3 LLM Provider配置参考', level=2)
add_table(doc,
    ['Provider', 'API_BASE', '推荐模型', '是否需要API_KEY'],
    [
        ['OpenAI', 'https://api.openai.com/v1', 'gpt-4o-mini / gpt-4o', '是'],
        ['DeepSeek', 'https://api.deepseek.com/v1', 'deepseek-chat / deepseek-reasoner', '是'],
        ['Ollama', 'http://localhost:11434/v1', 'qwen2.5:7b / llama3.1:8b', '否'],
        ['vLLM', 'http://localhost:8000/v1', '取决于加载的模型', '否'],
        ['LocalAI', 'http://localhost:8080/v1', '取决于加载的模型', '否'],
    ]
)

doc.add_heading('4. 数据扩展指南', level=1)
doc.add_paragraph('知识库数据存储在 backend/data/ 目录下的YAML文件中，可直接编辑扩展。')

doc.add_heading('4.1 添加新机器人方案', level=2)
doc.add_paragraph('编辑 backend/data/market_robots.yaml，在 robots 列表中添加新条目：')
doc.add_paragraph(
    '- name: 机器人名称\n'
    '  category: 分类（家用服务/户外服务/商用服务等）\n'
    '  description: 描述\n'
    '  price_range: 价格区间\n'
    '  diy_feasibility:\n'
    '    level: DIY难度（中等/中高/高/很高）\n'
    '    min_budget: 最低预算\n'
    '  tech_framework: 技术框架（可选）\n'
    '  detailed_bom: 物料清单（可选）\n'
    '  implementation_path: 实现路径（可选）\n'
    '  recommended_stack: 推荐技术栈\n'
    '  resource_links: 资源链接（可选）\n'
    '  open_source: 开源项目（可选）\n'
    '  attachments: 配套资料（可选）',
    style='No Spacing'
)

doc.add_heading('4.2 添加新算法', level=2)
doc.add_paragraph('在 backend/data/algorithms/ 目录下对应类别的YAML文件中添加：')
doc.add_paragraph(
    '- name: 算法名称\n'
    '  category: 类别（运动学/感知/步态规划/平衡控制/控制）\n'
    '  level: 级别（basic/intermediate/advanced）\n'
    '  description: 描述\n'
    '  applicable_robots: 适用机器人类型列表\n'
    '  compute_requirement: 计算需求\n'
    '  dependencies: 依赖库',
    style='No Spacing'
)

doc.add_heading('4.3 添加新硬件', level=2)
doc.add_paragraph('在 backend/data/hardware/ 目录下对应类别的YAML文件中添加：')
doc.add_paragraph(
    '- name: 硬件名称\n'
    '  category: 类别（sensors/actuators/controllers/power_communication/robot_kits）\n'
    '  type: 类型\n'
    '  price: 价格\n'
    '  specs: 规格参数\n'
    '  compatible_robots: 兼容机器人类型',
    style='No Spacing'
)

doc.add_heading('4.4 添加新技术栈', level=2)
doc.add_paragraph('在 backend/data/tech_stack/ 目录下对应类别的YAML文件中添加：')
doc.add_paragraph(
    '- name: 技术名称\n'
    '  category: 类别（运动控制/视觉/SLAM导航/仿真/语音LLM/部署）\n'
    '  description: 描述\n'
    '  tags: 标签列表\n'
    '  compatibility: 兼容平台\n'
    '  links: 相关链接',
    style='No Spacing'
)

doc.add_heading('5. API接口文档', level=1)
doc.add_paragraph('后端API基础路径：http://localhost:9500（开发模式通过Vite代理为 /api）')

doc.add_heading('5.1 健康检查', level=2)
add_table(doc,
    ['方法', '路径', '说明'],
    [
        ['GET', '/health', '系统健康检查'],
    ]
)

doc.add_heading('5.2 需求采集 API', level=2)
add_table(doc,
    ['方法', '路径', '说明'],
    [
        ['POST', '/api/requirement/session', '创建需求采集会话'],
        ['GET', '/api/requirement/session/{id}/round/{num}', '获取指定轮次问题'],
        ['POST', '/api/requirement/session/{id}/round/{num}', '提交指定轮次答案'],
        ['GET', '/api/requirement/session/{id}', '获取会话状态'],
        ['POST', '/api/requirement/session/{id}/build', '从会话构建需求对象'],
        ['GET', '/api/requirement/robot-types', '获取机器人类型选项'],
        ['POST', '/api/requirement/hardware-algorithm', '处理硬件算法适配请求'],
    ]
)

doc.add_heading('5.3 方案生成 API', level=2)
add_table(doc,
    ['方法', '路径', '说明'],
    [
        ['POST', '/api/solution/generate', '根据需求生成技术方案'],
    ]
)

doc.add_heading('5.4 算法库 API', level=2)
add_table(doc,
    ['方法', '路径', '参数', '说明'],
    [
        ['GET', '/api/algorithm/list', 'category/level/robot_type（可选）', '获取算法列表'],
        ['GET', '/api/algorithm/search', 'keyword（必填）', '搜索算法'],
        ['GET', '/api/algorithm/detail', 'name（必填）', '获取算法详情'],
        ['POST', '/api/algorithm/adapt', '-', '根据硬件适配算法'],
    ]
)

doc.add_heading('5.5 硬件方案 API', level=2)
add_table(doc,
    ['方法', '路径', '参数', '说明'],
    [
        ['GET', '/api/hardware/list', 'category/robot_type（可选）', '获取硬件列表'],
        ['GET', '/api/hardware/categories', '-', '获取硬件分类'],
        ['GET', '/api/hardware/search', 'keyword（必填）', '搜索硬件'],
        ['GET', '/api/hardware/kits', 'robot_type/level（可选）', '获取完整方案'],
        ['GET', '/api/hardware/budget', 'budget_min/budget_max（必填）, robot_type（可选）', '按预算筛选'],
        ['GET', '/api/hardware/item', 'name（必填）', '获取硬件详情'],
        ['GET', '/api/hardware/kit', 'name（必填）', '获取方案详情'],
        ['GET', '/api/hardware/tech-stack', 'category（可选）', '获取技术栈列表'],
    ]
)

doc.add_heading('5.6 矛盾检测 API', level=2)
add_table(doc,
    ['方法', '路径', '说明'],
    [
        ['POST', '/api/validator/requirement', '检测需求中的矛盾'],
        ['POST', '/api/validator/hardware-algorithm', '检测硬件与算法需求的矛盾'],
    ]
)

doc.add_heading('5.7 市场方案 API', level=2)
add_table(doc,
    ['方法', '路径', '参数', '说明'],
    [
        ['GET', '/api/market/robots', 'category（可选）', '获取市场机器人列表'],
        ['GET', '/api/market/robots/{name}', '-', '获取指定机器人详情'],
        ['GET', '/api/market/categories', '-', '获取机器人分类列表'],
        ['GET', '/api/market/frontier', '-', '获取前沿技术内容'],
        ['GET', '/api/market/basics', '-', '获取入门基础知识'],
    ]
)

doc.add_heading('6. 外部系统对接', level=1)

doc.add_heading('6.1 LLM API配置', level=2)
doc.add_paragraph(
    '系统支持通过标准OpenAI API格式接入各种LLM服务。配置步骤：\n'
    '1. 在 .env 文件中设置 APP_LLM_ENABLED=true\n'
    '2. 设置 APP_LLM_API_BASE 为对应服务的API地址\n'
    '3. 设置 APP_LLM_API_KEY（本地模型如Ollama无需设置）\n'
    '4. 设置 APP_LLM_MODEL 为对应模型名称\n'
    '5. 重启后端服务'
)

doc.add_heading('6.2 ROS2集成建议', level=2)
doc.add_paragraph(
    '系统当前不直接集成ROS2，但可通过以下方式对接：\n'
    '- 方案生成后，将技术方案中的ROS2包列表导出，用于rosdep安装\n'
    '- 通过HTTP API将需求/方案数据传递给ROS2节点\n'
    '- 使用rosbridge_suite将ROS2话题/服务暴露为WebSocket，前端直接订阅'
)

doc.add_heading('7. 与其他智能体互动', level=1)

doc.add_heading('7.1 MAVLink集成（无人机/无人车）', level=2)
doc.add_paragraph(
    '对于无人机和无人车方案，可通过MAVLink协议与飞控系统交互：\n'
    '- 使用pymavlink库通过串口/UDP连接飞控\n'
    '- 通过HTTP API将MAVLink数据传递给本系统进行方案优化\n'
    '- 系统生成的方案可包含ArduPilot/PX4配置建议'
)

doc.add_heading('7.2 ROS2 DDS通信', level=2)
doc.add_paragraph(
    '通过ROS2的DDS中间件实现与其他机器人智能体的数据交换：\n'
    '- 使用ros2cli工具发布/订阅话题\n'
    '- 通过rosbridge_server提供WebSocket接口\n'
    '- 本系统的方案数据可序列化为ROS2消息格式'
)

doc.add_heading('7.3 HTTP API对接', level=2)
doc.add_paragraph(
    '最简单的集成方式，其他智能体通过HTTP请求访问本系统：\n'
    '- 调用 /api/solution/generate 生成方案\n'
    '- 调用 /api/algorithm/adapt 适配算法\n'
    '- 调用 /api/hardware/budget 按预算筛选\n'
    '- 调用 /api/validator/requirement 检测矛盾\n'
    '所有API返回JSON格式数据，便于其他系统解析使用'
)

doc.add_heading('8. Docker部署指南', level=1)

doc.add_heading('8.1 快速部署', level=2)
doc.add_paragraph(
    '1. 确保已安装Docker和Docker Compose\n'
    '2. 在项目根目录执行：docker-compose up -d\n'
    '3. 访问 http://localhost:3000 使用系统\n'
    '4. 后端API地址：http://localhost:8000'
)

doc.add_heading('8.2 配置LLM（Docker环境）', level=2)
doc.add_paragraph(
    '在docker-compose.yaml中添加环境变量：\n'
    'environment:\n'
    '  - APP_LLM_ENABLED=true\n'
    '  - APP_LLM_API_KEY=your-api-key\n'
    '  - APP_LLM_API_BASE=https://api.openai.com/v1\n'
    '  - APP_LLM_MODEL=gpt-4o-mini'
)

doc.add_heading('8.3 数据持久化', level=2)
doc.add_paragraph(
    'Docker部署时，YAML数据文件通过volume挂载：\n'
    'volumes:\n'
    '  - ./backend/data:/app/data\n'
    '修改本地data目录下的YAML文件即可更新知识库数据，无需重新构建镜像。'
)

doc.add_heading('9. 常见问题FAQ', level=1)

faqs = [
    ('Q: 如何启用LLM功能？',
     'A: 在backend/.env文件中设置APP_LLM_ENABLED=true，并配置对应的API密钥和地址。'
        '如果使用Ollama等本地模型，只需设置API地址即可，无需API密钥。'),
    ('Q: LLM调用失败怎么办？',
     'A: 系统默认开启回退机制（APP_LLM_FALLBACK_ENABLED=true），LLM调用失败时会自动使用预设规则。'
        '检查API密钥、网络连接和模型名称是否正确。'),
    ('Q: 如何添加新的机器人类型？',
     'A: 编辑backend/data/market_robots.yaml，按照现有格式添加新条目。'
        '同时在schemas/requirement.py中添加对应的RobotType枚举值。'),
    ('Q: 前端如何连接后端？',
     'A: 开发模式下，Vite配置了代理（vite.config.ts），将/api请求转发到http://localhost:9500。'
        '生产模式下，Nginx配置了反向代理。'),
    ('Q: Docker部署后无法访问？',
     'A: 检查端口是否被占用（前端3000/后端8000），查看容器日志：docker-compose logs。'
        '确保backend容器健康检查通过。'),
    ('Q: 如何运行测试？',
     'A: 后端测试：cd backend && pytest tests/ -v\n'
        '前端E2E测试：cd frontend && npx playwright test --reporter=list\n'
        '确保前后端服务已启动。'),
    ('Q: 知识库数据更新后需要重启吗？',
     'A: 是的，当前版本需要重启后端服务才能加载新的YAML数据。'
        'Docker环境下执行：docker-compose restart backend'),
]

for q, a in faqs:
    doc.add_paragraph(q, style='List Bullet')
    doc.add_paragraph(a)

doc.save(r'e:\个人\机器人应用\通用机器人视觉运动系统\docs\项目说明书.docx')
print("项目说明书.docx generated successfully")
