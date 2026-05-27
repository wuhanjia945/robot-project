from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)

doc.add_heading('一起动手来手搓机器人 - 项目版本记录', 0)

versions = [
    {
        "version": "v1.0.0",
        "date": "2026-05-27",
        "changes": [
            "正式版本发布",
            "后端：FastAPI框架，6大核心模块（需求采集/方案生成/算法适配/硬件方案/矛盾检测/知识库）",
            "前端：React+TypeScript+Vite，8个页面（首页/市场方案/需求采集/方案生成/算法供给/硬件方案/仿真学习/反馈）",
            "数据：YAML结构化知识库（10种机器人方案/22+算法/26+硬件/6种仿真软件/5项前沿技术）",
            "LLM接入：预设优先+LLM兜底，支持OpenAI/DeepSeek/Ollama/vLLM/LocalAI",
            "测试：50个后端单元测试 + 12个Playwright前端E2E测试（全部通过）",
            "部署：Docker Compose容器化部署",
            "文档：项目版本记录/项目说明书/测试结果报告",
        ]
    },
    {
        "version": "v0.3.0",
        "date": "2026-05-26",
        "changes": [
            "新增LLM接入功能（预设优先+LLM兜底架构）",
            "支持云端API（OpenAI/DeepSeek）和本地模型（Ollama/vLLM/LocalAI）",
            "添加4个专业提示词模板（机器人专家/方案生成/需求分析/算法推荐）",
            "方案生成器支持LLM兜底：预算超限/特殊功能/未知类型时自动调用LLM",
            "需求引擎支持LLM自然语言分析",
            "硬件方案支持LLM算法推荐",
        ]
    },
    {
        "version": "v0.2.0",
        "date": "2026-05-26",
        "changes": [
            "市场方案：添加详细技术框架+实现路径+实际BOM清单",
            "每个方案添加附件（教程+配套源代码），链接全部替换为真实开源项目",
            "新增仿真学习模块（Gazebo/MuJoCo/Isaac Sim/PyBullet/Webots/Genie Sim）",
            "新增用户反馈收集入口",
            "修复MarketPage详情视图按钮切换Bug",
            "补充宇树科技/智元机器人技术栈和开源项目",
            "添加5项前沿技术内容（具身智能/RL运动策略/大模型+机器人/Sim2Real/VLA）",
        ]
    },
    {
        "version": "v0.1.0",
        "date": "2026-05-25",
        "changes": [
            "项目初始化",
            "后端FastAPI框架搭建，6大核心模块实现",
            "前端React+TypeScript+Vite搭建",
            "YAML知识库数据填充",
            "50个后端单元测试通过",
            "Docker Compose容器化部署配置",
        ]
    },
]

for v in versions:
    doc.add_heading(f'{v["version"]} ({v["date"]})', level=1)
    for change in v["changes"]:
        doc.add_paragraph(change, style='List Bullet')

doc.save(r'e:\个人\机器人应用\通用机器人视觉运动系统\docs\版本记录.docx')
print("版本记录.docx generated successfully")
